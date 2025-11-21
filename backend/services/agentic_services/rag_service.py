import os
import uuid
from datetime import datetime
from werkzeug.utils import secure_filename
from pypdf import PdfReader
from docx import Document as DocxDocument
import chromadb
from chromadb.config import Settings
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.tools import DuckDuckGoSearchRun

from models import db
from models.agentic_entities.document import Document
from models.agentic_entities.chat import ChatHistory
from services.auth_services.auth_service import AuthService
from config import Config

class RAGService:
    """LangChain Agentic RAG service for document management and chat"""
    
    def __init__(self):
        """Initialize RAG service"""
        if not Config.OPENAI_API_KEY:
            raise ValueError('OpenAI API key not configured')
        
        # Initialize OpenAI embeddings
        self.embeddings = OpenAIEmbeddings(
            model="text-embedding-3-small",
            openai_api_key=Config.OPENAI_API_KEY
        )
        
        # Initialize OpenAI model
        self.llm = ChatOpenAI(
            model=Config.OPENAI_MODEL,
            openai_api_key=Config.OPENAI_API_KEY,
            temperature=0.7
        )
        
        # Initialize Chroma client
        self.chroma_client = chromadb.PersistentClient(
            path=Config.CHROMA_DB_PATH
        )
        
        # Initialize web search tool
        self.search_tool = DuckDuckGoSearchRun()
    
    def upload_document(self, file, user_id):
        """
        Upload and process document for RAG
        
        Args:
            file: File object from request
            user_id: User ID
            
        Returns:
            dict: Document data
            
        Raises:
            ValueError: If file is invalid
        """
        # Validate file
        if not file:
            raise ValueError('No file provided')
        
        filename = secure_filename(file.filename)
        if not filename:
            raise ValueError('Invalid filename')
        
        # Check file extension
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        if ext not in Config.ALLOWED_EXTENSIONS:
            raise ValueError(f'File type not allowed. Allowed types: {Config.ALLOWED_EXTENSIONS}')
        
        # Save file
        file_id = str(uuid.uuid4())
        filepath = os.path.join(Config.DOCUMENTS_PATH, f"{file_id}_{filename}")
        file.save(filepath)
        
        # Get file size
        file_size = os.path.getsize(filepath)
        
        # Extract text content
        try:
            text_content = self._extract_text(filepath, ext)
        except Exception as e:
            os.remove(filepath)
            raise ValueError(f'Error extracting text: {str(e)}')
        
        # Create vector store for this document
        try:
            collection_name = f"doc_{file_id}"
            vector_store = self._create_vector_store(text_content, collection_name)
        except Exception as e:
            os.remove(filepath)
            raise ValueError(f'Error creating vector store: {str(e)}')
        
        # Save document to database
        document = Document(
            filename=filename,
            filepath=filepath,
            user_id=user_id,
            vector_store_id=collection_name,
            file_size=file_size
        )
        db.session.add(document)
        db.session.commit()
        
        return document.to_dict()
    
    def _extract_text(self, filepath, ext):
        """Extract text from document"""
        text = ""
        
        if ext == 'pdf':
            reader = PdfReader(filepath)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        
        elif ext in ['docx', 'doc']:
            doc = DocxDocument(filepath)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        
        elif ext in ['txt', 'md']:
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        
        else:
            raise ValueError(f'Unsupported file type: {ext}')
        
        if not text.strip():
            raise ValueError('No text content found in document')
        
        return text
    
    def _create_vector_store(self, text, collection_name):
        """Create vector store from text"""
        # Split text into chunks
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        chunks = text_splitter.split_text(text)
        
        # Create Chroma vector store
        vector_store = Chroma.from_texts(
            texts=chunks,
            embedding=self.embeddings,
            collection_name=collection_name,
            client=self.chroma_client
        )
        
        return vector_store
    
    def chat_with_documents(self, query, user_id, use_internet=False):
        """
        Chat with user's documents using RAG
        
        Args:
            query: User question
            user_id: User ID
            use_internet: Whether to use internet search
            
        Returns:
            dict: Response with answer and sources
        """
        # Get user's documents
        documents = Document.query.filter_by(user_id=user_id).all()
        
        if not documents:
            raise ValueError('No documents found. Please upload documents first.')
        
        # Combine all document collections
        all_docs = []
        for doc in documents:
            try:
                vector_store = Chroma(
                    collection_name=doc.vector_store_id,
                    embedding_function=self.embeddings,
                    client=self.chroma_client
                )
                # Retrieve relevant documents
                retrieved_docs = vector_store.similarity_search(query, k=3)
                all_docs.extend(retrieved_docs)
            except Exception as e:
                print(f"Error retrieving from {doc.filename}: {e}")
        
        if not all_docs:
            raise ValueError('Could not retrieve relevant information from documents')
        
        # Build context from documents
        context = "\n\n".join([doc.page_content for doc in all_docs])
        
        # Add internet search if enabled
        internet_info = ""
        if use_internet:
            try:
                search_results = self.search_tool.run(query)
                internet_info = f"\n\nInternet Search Results:\n{search_results}"
            except Exception as e:
                print(f"Internet search error: {e}")
        
        # Create prompt
        prompt_template = """You are a helpful AI assistant. Answer the question based on the provided context and your knowledge.
        
        Context from documents:
        {context}
        {internet_info}
        
        Question: {query}
        
        Provide a comprehensive and accurate answer. If the answer is not in the context, say so and provide general knowledge if appropriate."""
        
        prompt = prompt_template.format(
            context=context,
            internet_info=internet_info,
            query=query
        )
        
        # Generate response
        response = self.llm.invoke(prompt)
        answer = response.content
        
        # Save chat history
        chat_history = ChatHistory(
            user_id=user_id,
            message=query,
            response=answer,
            chat_type='rag',
            extra_metadata={
                'use_internet': use_internet,
                'num_sources': len(all_docs)
            }
        )
        db.session.add(chat_history)
        db.session.commit()
        
        return {
            'answer': answer,
            'sources': [
                {
                    'content': doc.page_content[:200] + '...',
                    'metadata': doc.metadata if hasattr(doc, 'metadata') else {}
                }
                for doc in all_docs[:3]
            ],
            'use_internet': use_internet
        }
    
    def get_user_documents(self, user_id):
        """Get all documents for a user"""
        documents = Document.query.filter_by(user_id=user_id).all()
        return [doc.to_dict() for doc in documents]
    
    def delete_document(self, document_id, user_id):
        """
        Delete a document
        
        Args:
            document_id: Document ID
            user_id: User ID
            
        Returns:
            dict: Success message
        """
        document = Document.query.filter_by(id=document_id, user_id=user_id).first()
        
        if not document:
            raise ValueError('Document not found')
        
        # Delete file
        try:
            if os.path.exists(document.filepath):
                os.remove(document.filepath)
        except Exception as e:
            print(f"Error deleting file: {e}")
        
        # Delete vector store collection
        try:
            self.chroma_client.delete_collection(document.vector_store_id)
        except Exception as e:
            print(f"Error deleting collection: {e}")
        
        # Delete from database
        db.session.delete(document)
        db.session.commit()
        
        return {'message': f'Document {document.filename} deleted successfully'}
